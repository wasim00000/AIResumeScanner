import logging
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def extract_text_from_pdf(pdf_path):
    """
    Extract text from a PDF file

    Args:
        pdf_path (str): Path to the PDF file

    Returns:
        str: Extracted text from the PDF
    """
    try:
        text = ""
        pdf_path = Path(pdf_path)
        with open(pdf_path, 'rb') as file:
            import PyPDF2

            pdf_reader = PyPDF2.PdfReader(file)
            # Extract text from each page
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text() or ""
                text += page_text + "\n"

        if not text.strip():
            logger.warning(f"No text extracted from PDF: {pdf_path}")
            return "No readable text found in PDF"

        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF {pdf_path}: {str(e)}")
        raise Exception(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_docx(docx_path):
    """
    Extract text from a DOCX file

    Args:
        docx_path (str): Path to the DOCX file

    Returns:
        str: Extracted text from the DOCX
    """
    try:
        try:
            import docx
        except ImportError as exc:
            raise ImportError("python-docx is required to extract text from DOCX files") from exc

        doc = docx.Document(str(Path(docx_path)))
        text = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            text.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)

        full_text = "\n".join(text)

        if not full_text.strip():
            logger.warning(f"No text extracted from DOCX: {docx_path}")
            return "No readable text found in DOCX"

        return full_text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {docx_path}: {str(e)}")
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")
