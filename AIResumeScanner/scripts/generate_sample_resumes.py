from __future__ import annotations

import json
import random
import sys
from datetime import datetime
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from database import save_analysis_result, save_job_description, save_resume  # noqa: E402
from nlp_processor import extract_entities, extract_skills, preprocess_text  # noqa: E402
from ranking_system import calculate_similarity  # noqa: E402
from resume_parser import extract_text_from_docx  # noqa: E402


SAMPLE_DIR = PROJECT_ROOT / "sample_resumes"
SAMPLE_DIR.mkdir(parents=True, exist_ok=True)


FIRST_NAMES = [
    "Ava", "Liam", "Noor", "Ethan", "Priya",
    "Mateo", "Sofia", "Arjun", "Chloe", "Daniel",
]

LAST_NAMES = ["Patel", "Chen", "Johnson", "Khan", "Rivera"]

COMPANIES = [
    "Northstar Analytics LLC",
    "Bluepeak Systems Inc",
    "Meridian Cloud Partners",
    "Vertex Labs Ltd",
    "Summit Digital Solutions",
    "Orchard Tech Group",
    "Crestline Software Corp",
    "Pioneer Automation LLC",
    "Horizon Data Services",
    "Lighthouse Engineering Ltd",
]

UNIVERSITIES = [
    "University of California, Irvine",
    "Georgia Institute of Technology",
    "University of Texas at Austin",
    "University of Washington",
    "Arizona State University",
    "Purdue University",
]

PERSONAS = [
    {
        "role": "Senior Backend Engineer",
        "summary": "Builds resilient APIs, background jobs, and integration services with Python, Flask, FastAPI, PostgreSQL, and AWS.",
        "skills": ["Python", "Flask", "FastAPI", "PostgreSQL", "AWS", "Docker", "CI/CD", "Git", "Redis", "REST", "Node.js", "SQL"],
        "certifications": ["AWS Certified Solutions Architect"],
        "education": "Bachelor of Science in Computer Science",
    },
    {
        "role": "Full Stack Engineer",
        "summary": "Designs customer-facing products with React, TypeScript, Node.js, and PostgreSQL while keeping delivery predictable.",
        "skills": ["JavaScript", "TypeScript", "React", "Node.js", "Express", "MongoDB", "SQL", "Git", "Jira", "REST", "HTML", "CSS"],
        "certifications": ["Scrum Foundation Professional Certificate"],
        "education": "Bachelor of Engineering in Information Technology",
    },
    {
        "role": "Data Scientist",
        "summary": "Turns raw datasets into reliable models and dashboards using Python, pandas, NumPy, scikit-learn, and SQL.",
        "skills": ["Python", "Pandas", "NumPy", "scikit-learn", "SQL", "Tableau", "Statistics", "Machine Learning", "Data Analysis", "Visualization", "Git", "A/B Testing"],
        "certifications": ["Google Data Analytics Professional Certificate"],
        "education": "Master of Science in Data Science",
    },
    {
        "role": "DevOps Engineer",
        "summary": "Keeps delivery pipelines stable with Docker, Kubernetes, Terraform, Jenkins, cloud infrastructure, and observability tooling.",
        "skills": ["Docker", "Kubernetes", "Terraform", "Jenkins", "AWS", "Linux", "CI/CD", "Git", "Shell", "Monitoring", "Python", "Automation"],
        "certifications": ["AWS Certified DevOps Engineer"],
        "education": "Bachelor of Science in Software Engineering",
    },
    {
        "role": "Product Analyst",
        "summary": "Connects business questions to product metrics, customer insights, and decision support with SQL, Tableau, and clear communication.",
        "skills": ["SQL", "Excel", "Tableau", "Power BI", "Communication", "Stakeholder Management", "Data Analysis", "Reporting", "A/B Testing", "Presentation", "Python", "Leadership"],
        "certifications": ["Certified Analytics Professional"],
        "education": "Bachelor of Arts in Economics",
    },
]


SAMPLE_JOB_DESCRIPTION = """
Senior Full-Stack Data Platform Engineer

We are looking for a professional who can build reliable products and internal tools across Python, React, Node.js, SQL, PostgreSQL, AWS, Docker, Kubernetes, and CI/CD pipelines.

Requirements:
- 5+ years of experience in software engineering or data engineering
- Strong knowledge of Python, JavaScript, TypeScript, React, Node.js, and SQL
- Experience with cloud platforms, monitoring, and automation
- Familiarity with Git, Docker, Kubernetes, Terraform, and REST APIs
- Strong communication, teamwork, and problem solving skills

Preferred:
- FastAPI, Flask, Redis, and testing frameworks
- Background in analytics, dashboards, or machine learning
- Bachelor degree in Computer Science or equivalent
""".strip()


def make_slug(name: str) -> str:
    return name.lower().replace(" ", "-")


def build_profiles() -> list[dict]:
    rng = random.Random(42)
    profiles: list[dict] = []
    name_pairs = [(first, last) for last in LAST_NAMES for first in FIRST_NAMES]

    for index, (first, last) in enumerate(name_pairs):
        persona = PERSONAS[index % len(PERSONAS)]
        years = 3 + (index % 9)
        primary_company = COMPANIES[index % len(COMPANIES)]
        secondary_company = COMPANIES[(index + 4) % len(COMPANIES)]
        university = UNIVERSITIES[index % len(UNIVERSITIES)]
        skills = persona["skills"][:]
        rng.shuffle(skills)
        selected_skills = skills[:10]
        selected_skills.extend(["Communication", "Teamwork"])
        selected_skills = list(dict.fromkeys(selected_skills))

        profiles.append({
            "name": f"{first} {last}",
            "slug": make_slug(f"{first}-{last}"),
            "email": f"{make_slug(f'{first}.{last}')}@example.com",
            "phone": f"555-01{index:02d}",
            "linkedin": f"https://linkedin.com/in/{make_slug(f'{first}{last}')}",
            "github": f"https://github.com/{make_slug(f'{first}{last}')}",
            "location": f"{['Austin', 'Seattle', 'Denver', 'Chicago', 'Boston'][index % 5]}, USA",
            "role": persona["role"],
            "summary": persona["summary"],
            "skills": selected_skills,
            "certifications": persona["certifications"],
            "education": f"{persona['education']}, {university}",
            "experience_years": years,
            "companies": [primary_company, secondary_company],
        })

    return profiles


def write_resume_docx(profile: dict, output_path: Path) -> None:
    doc = Document()
    current_year = datetime.now().year

    doc.add_heading(profile["name"], level=0)
    doc.add_paragraph(profile["role"])
    doc.add_paragraph(profile["summary"])
    doc.add_paragraph(f"Email: {profile['email']}")
    doc.add_paragraph(f"Phone: {profile['phone']}")
    doc.add_paragraph(f"LinkedIn: {profile['linkedin']}")
    doc.add_paragraph(f"GitHub: {profile['github']}")
    doc.add_paragraph(f"Location: {profile['location']}")

    doc.add_heading("Professional Skills", level=1)
    doc.add_paragraph(", ".join(profile["skills"]))

    doc.add_heading("Professional Experience", level=1)
    primary_company, secondary_company = profile["companies"]
    years = profile["experience_years"]
    doc.add_paragraph(f"Senior Contributor | {primary_company} | {current_year - years}-{current_year}")
    doc.add_paragraph(
        f"Led delivery of internal tools, APIs, and reporting workflows with {years} years experience across cross-functional teams."
    )
    doc.add_paragraph(f"Lead Specialist | {secondary_company} | {current_year - min(years, 4)}-{current_year - 1}")
    doc.add_paragraph(
        "Partnered with engineering, product, and operations teams to improve reliability, automation, and stakeholder visibility."
    )

    doc.add_heading("Education", level=1)
    doc.add_paragraph(f"{profile['education']}")

    doc.add_heading("Certifications", level=1)
    for certification in profile["certifications"]:
        doc.add_paragraph(certification)

    doc.add_heading("Selected Projects", level=1)
    doc.add_paragraph(
        f"Built a production dashboard using Python, React, SQL, and AWS to monitor data quality and surface operational issues."
    )
    doc.add_paragraph(
        f"Created a CI/CD workflow with Docker, Git, and automated checks to improve release confidence and deployment speed."
    )

    doc.add_heading("Highlights", level=1)
    doc.add_paragraph(
        f"Known for {profile['role'].lower()} execution, analytical thinking, communication, and reliable follow-through."
    )

    doc.save(output_path)


def seed_database(resume_files: list[Path], manifest: list[dict]) -> dict:
    preprocessed_jd = preprocess_text(SAMPLE_JOB_DESCRIPTION)
    job_skills = extract_skills(preprocessed_jd)
    job_id = save_job_description(SAMPLE_JOB_DESCRIPTION, job_skills)

    seeded = []

    for file_path, entry in zip(resume_files, manifest):
        text = extract_text_from_docx(file_path)
        preprocessed_text = preprocess_text(text)
        resume_skills = extract_skills(preprocessed_text)
        entities = extract_entities(text)
        candidate_name = entry["name"]
        for entity_text, entity_type in entities:
            if entity_type == "PERSON":
                candidate_name = entity_text
                break

        score = calculate_similarity(preprocessed_jd, preprocessed_text, job_skills, resume_skills)
        matching_skills = sorted(set(job_skills).intersection(set(resume_skills)))
        resume_id = save_resume(file_path.name, candidate_name, text, resume_skills)
        analysis_id = save_analysis_result(job_id, resume_id, score, matching_skills)

        seeded.append({
            "resume_id": resume_id,
            "analysis_id": analysis_id,
            "candidate_name": candidate_name,
            "filename": file_path.name,
            "match_percentage": round(score * 100, 2),
        })

    return {
        "job_id": job_id,
        "job_skills": job_skills,
        "seeded_resumes": seeded,
    }


def main() -> None:
    profiles = build_profiles()
    manifest: list[dict] = []
    resume_files: list[Path] = []

    for index, profile in enumerate(profiles, start=1):
        filename = f"{index:02d}_{profile['slug']}_resume.docx"
        output_path = SAMPLE_DIR / filename
        write_resume_docx(profile, output_path)
        resume_files.append(output_path)
        manifest.append({
            "file": filename,
            "name": profile["name"],
            "role": profile["role"],
            "skills": profile["skills"],
            "location": profile["location"],
        })

    seed_info = seed_database(resume_files, manifest)

    generated_at = datetime.now().isoformat()
    (SAMPLE_DIR / "manifest.json").write_text(
        json.dumps(
            {
                "generated_at": generated_at,
                "job_id": seed_info["job_id"],
                "job_skills": seed_info["job_skills"],
                "resumes": manifest,
                "seeded_resumes": seed_info["seeded_resumes"],
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    (SAMPLE_DIR / "README.md").write_text(
        "\n".join(
            [
                "# Sample Resumes",
                "",
                "This folder contains 50 generated DOCX resumes for local testing and documentation.",
                "",
                f"Generated at: {generated_at}",
                "",
                "Run `python scripts/generate_sample_resumes.py` from the AIResumeScanner directory to regenerate the sample set.",
            ]
        )
        + "\n",
        encoding="utf-8",
    )

    print(f"Generated {len(resume_files)} resumes in {SAMPLE_DIR}")
    print(f"Seeded job description ID: {seed_info['job_id']}")
    print(f"Seeded resume records: {len(seed_info['seeded_resumes'])}")


if __name__ == "__main__":
    main()
