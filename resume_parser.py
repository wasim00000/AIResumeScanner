
import PyPDF2
import docx
import pytesseract
from pdf2image import convert_from_path
import io
import os
import logging
import numpy as np
from doctr.io import DocumentFile
from doctr.models import ocr_predictor

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Initialize OCR model
predictor = ocr_predictor(det_arch='db_resnet50', reco_arch='crnn_vgg16_bn', pretrained=True)

def extract_text_from_pdf(pdf_path):
    """
    Extract text from a PDF file using multiple methods
    """
    try:
        text = ""
        # Try normal PDF extraction first
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
        
        # If no text found, use OCR
        if not text.strip():
            # Convert PDF to images
            images = convert_from_path(pdf_path)
            for image in images:
                # Use doctr for better OCR
                doc = DocumentFile.from_images(np.array(image))
                result = predictor(doc)
                text += result.render() + "\n"
                
                # Fallback to tesseract if needed
                if not text.strip():
                    text += pytesseract.image_to_string(image) + "\n"
        
        if not text.strip():
            logger.warning(f"No text extracted from PDF: {pdf_path}")
            return "No readable text found in PDF"
            
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF {pdf_path}: {str(e)}")
        raise Exception(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_docx(docx_path):
    """
    Extract text from a DOCX file with enhanced extraction
    """
    try:
        doc = docx.Document(docx_path)
        text = []
        
        # Extract text from paragraphs with formatting
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                text.append(f"\n## {para.text}\n")
            else:
                text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            text.append("\nTable content:")
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                text.append(" | ".join(row_text))
        
        full_text = "\n".join(text)
        
        if not full_text.strip():
            logger.warning(f"No text extracted from DOCX: {docx_path}")
            return "No readable text found in DOCX"
            
        return full_text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {docx_path}: {str(e)}")
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")
